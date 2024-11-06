# aws_translator_app/services/document_service.py

"""
Document Service Module
=======================

Este módulo fornece serviços para importar e exportar textos a partir e para
diferentes formatos de documentos. Suporta os seguintes formatos:

- **Importação**:
    - PDF (`.pdf`)
    - DOCX (`.docx`)
    - EPUB (`.epub`)
    - TXT (`.txt`)
- **Exportação**:
    - PDF (`.pdf`)
    - DOCX (`.docx`)
    - TXT (`.txt`)

Classes:
    DocumentService: Classe responsável pela importação e exportação de documentos.

Dependências:
    - PyPDF2: biblioteca para manipulação de arquivos PDF.
    - python-docx: biblioteca para manipulação de arquivos DOCX.
    - EbookLib: biblioteca para manipulação de arquivos EPUB.
    - reportlab: biblioteca para geração de PDFs.
    - typing: biblioteca padrão para anotações de tipos.

Exemplo de Uso:
    >>> from translation_app.services.document_service import DocumentService
    >>> doc_service = DocumentService()

    # Importar um documento
    >>> texto = doc_service.import_document(file_object)

    # Exportar um documento
    >>> file_path = doc_service.export_document(text, metrics_original, metrics_simplified, format)
"""

from typing import Optional
import os

import PyPDF2  # Para PDFs
from docx import Document  # Para DOCX
from ebooklib import epub  # Para EPUB

from reportlab.lib.pagesizes import letter  # Para exportar PDFs
from reportlab.pdfgen import canvas


class DocumentService:
    """
    Serviço para importar e exportar textos a partir e para diferentes formatos de documentos.

    Esta classe fornece métodos para importar textos de arquivos PDF, DOCX, EPUB e TXT, bem
    como exportar textos para arquivos PDF, DOCX e TXT. Utiliza diversas bibliotecas para
    manipulação de diferentes formatos de arquivos.

    Métodos:
        import_document(file) ⇾ Optional[str]:
            Importa texto de um arquivo de documento.

        export_document(text: str, metrics_original: dict, metrics_simplified: dict, format: str) ⇾ str:
            Exporta texto e métricas para um arquivo de documento e retorna o caminho do arquivo gerado.
    """

    def import_document(self, file) -> Optional[str]:
        """
        Importa texto de um arquivo de documento.

        Este metodo determina o tipo de arquivo com base na extensão e utiliza o metodo
        apropriado para extrair o texto.

        Parâmetros:
            file: Arquivo enviado (UploadedFile).

        Retorna:
            Optional[str]: O texto extraído do documento ou `None` se não for possível extrair.

        Exceções:
            - ValueError: se o formato do arquivo não for suportado.
            - Exception: Se ocorrer um erro durante a importação do documento.
        """
        _, ext = os.path.splitext(file.name)
        ext = ext.lower()
        if ext == '.pdf':
            return self._import_pdf(file)
        elif ext == '.docx':
            return self._import_docx(file)
        elif ext == '.epub':
            return self._import_epub(file)
        elif ext == '.txt':
            return self._import_txt(file)
        else:
            raise ValueError(f"Formato de arquivo não suportado: {ext}")

    def export_document(self, text: str, metrics_original: dict, metrics_simplified: dict, format: str) -> str:
        """
        Exporta texto e métricas para um arquivo de documento.

        Este metodo determina o formato de exportação com base no parâmetro `format`
        e utiliza o metodo apropriado para salvar o texto e métricas no formato desejado.

        Parâmetros:
            text (str): O texto a ser exportado.
            metrics_original (dict): Métricas do texto original.
            metrics_simplified (dict): Métricas do texto simplificado.
            format (str): Formato de exportação desejado (`'pdf'`, `'docx'`, `'txt'`).

        Retorna:
            str: O caminho do arquivo exportado.

        Exceções:
            - ValueError: se o formato de exportação não for suportado.
            - Exception: Se ocorrer um erro durante a exportação do documento.
        """
        format = format.lower()
        # Gerar um caminho temporário para o arquivo
        file_path = f'/tmp/output.{format}'
        if format == 'pdf':
            self._export_pdf(text, file_path, metrics_original, metrics_simplified)
        elif format == 'docx':
            self._export_docx(text, file_path, metrics_original, metrics_simplified)
        elif format == 'txt':
            self._export_txt(text, file_path, metrics_original, metrics_simplified)
        else:
            raise ValueError(f"Formato de exportação não suportado: {format}")
        return file_path

    @staticmethod
    def _import_pdf(file) -> str:
        """
        Importa texto de um arquivo PDF.

        Utiliza a biblioteca PyPDF2 para extrair o texto de cada página do PDF.

        Parâmetros:
            file: Arquivo PDF enviado (UploadedFile).

        Retorna:
            str: O texto extraído do PDF.

        Exceções:
            - Exception: Se ocorrer um erro durante a leitura do PDF.
        """
        try:
            reader = PyPDF2.PdfReader(file)
            text = ''
            for page in reader.pages:
                extracted_text = page.extract_text()
                if extracted_text:
                    text += extracted_text + '\n'
            return text.strip()
        except Exception as e:
            raise Exception(f"Erro ao importar PDF: {str(e)}")

    @staticmethod
    def _import_docx(file) -> str:
        """
        Importa texto de um arquivo DOCX.

        Utiliza a biblioteca python-docx para extrair o texto de cada parágrafo do DOCX.

        Parâmetros:
            file: Arquivo DOCX enviado (UploadedFile).

        Retorna:
            str: O texto extraído do DOCX.

        Exceções:
            - Exception: Se ocorrer um erro durante a leitura do DOCX.
        """
        try:
            doc = Document(file)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return text.strip()
        except Exception as e:
            raise Exception(f"Erro ao importar DOCX: {str(e)}")

    @staticmethod
    def _import_epub(file) -> str:
        """
        Importa texto de um arquivo EPUB.

        Utiliza a biblioteca EbookLib para extrair o conteúdo textual dos documentos do EPUB.

        Parâmetros:
            file: Arquivo EPUB enviado (UploadedFile).

        Retorna:
            str: O texto extraído do EPUB.

        Exceções:
            - Exception: Se ocorrer um erro durante a leitura do EPUB.
        """
        try:
            book = epub.read_epub(file)
            text = ''
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    content = item.get_content()
                    text += content.decode('utf-8') + '\n'
            return text.strip()
        except Exception as e:
            raise Exception(f"Erro ao importar EPUB: {str(e)}")

    @staticmethod
    def _import_txt(file) -> str:
        """
        Importa texto de um arquivo TXT.

        Abre o arquivo de texto e lê todos o seu conteúdo.

        Parâmetros:
            file: Arquivo TXT enviado (UploadedFile).

        Retorna:
            str: O texto extraído do TXT.

        Exceções:
            - Exception: Se ocorrer um erro durante a leitura do TXT.
        """
        try:
            return file.read().decode('utf-8').strip()
        except Exception as e:
            raise Exception(f"Erro ao importar TXT: {str(e)}")

    @staticmethod
    def _export_pdf(text: str, file_path: str, metrics_original: dict = None, metrics_simplified: dict = None) -> None:
        """
        Exporta texto e métricas para um arquivo PDF.

        Utiliza a biblioteca ReportLab para gerar um PDF a partir do texto fornecido,
        cuidando da formatação e quebra de linhas conforme necessário, e inclui as métricas.

        Parâmetros:
            text (str): O texto a ser exportado para o PDF.
            file_path (str): Caminho onde o arquivo PDF será salvo.
            metrics_original (dict): Métricas do texto original.
            metrics_simplified (dict): Métricas do texto simplificado.

        Retorna:
            None

        Exceções:
            - Exception: Se ocorrer um erro durante a criação do PDF.
        """
        try:
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter

            # Configurações do texto
            text_object = c.beginText(50, height - 50)
            text_object.setFont("Helvetica-Bold", 14)
            text_object.textLine("Texto Simplificado e Traduzido:")
            text_object.setFont("Helvetica", 12)
            text_object.textLine("")

            # Adicionar o texto
            for line in text.split('\n'):
                words = line.split(' ')
                line_buffer = ""

                for word in words:
                    if c.stringWidth(line_buffer + word, "Helvetica", 12) < (width - 100):
                        line_buffer += word + " "
                    else:
                        text_object.textLine(line_buffer.strip())
                        line_buffer = word + " "

                        if text_object.getY() <= 50:
                            c.drawText(text_object)
                            c.showPage()
                            text_object = c.beginText(50, height - 50)
                            text_object.setFont("Helvetica", 12)

                if line_buffer:
                    text_object.textLine(line_buffer.strip())

                    if text_object.getY() <= 50:
                        c.drawText(text_object)
                        c.showPage()
                        text_object = c.beginText(50, height - 50)
                        text_object.setFont("Helvetica", 12)

            text_object.textLine("")

            if metrics_original and metrics_simplified:
                metric_names = {
                    'flesch_reading_ease': 'Índice de Flesch Reading Ease',
                    'flesch_kincaid_grade': 'Grau de Flesch-Kincaid',
                    'smog_index': 'Índice SMOG',
                    'coleman_liau_index': 'Índice de Coleman-Liau',
                    'automated_readability_index': 'Índice ARI',
                    'dale_chall_readability_score': 'Pontuação de Dale-Chall'
                }

                # Métricas do texto original
                text_object.setFont("Helvetica-Bold", 14)
                text_object.textLine("Métricas do Texto Original:")
                text_object.setFont("Helvetica", 12)
                text_object.textLine("")

                for key, value in metrics_original.items():
                    metric_name = metric_names.get(key, key)
                    text_object.textLine(f"{metric_name}: {value:.2f}")

                    if text_object.getY() <= 50:
                        c.drawText(text_object)
                        c.showPage()
                        text_object = c.beginText(50, height - 50)
                        text_object.setFont("Helvetica", 12)

                text_object.textLine("")

                # Métricas do texto simplificado
                text_object.setFont("Helvetica-Bold", 14)
                text_object.textLine("Métricas do Texto Simplificado:")
                text_object.setFont("Helvetica", 12)
                text_object.textLine("")

                for key, value in metrics_simplified.items():
                    metric_name = metric_names.get(key, key)
                    text_object.textLine(f"{metric_name}: {value:.2f}")

                    if text_object.getY() <= 50:
                        c.drawText(text_object)
                        c.showPage()
                        text_object = c.beginText(50, height - 50)
                        text_object.setFont("Helvetica", 12)

            c.drawText(text_object)
            c.save()
        except Exception as e:
            raise Exception(f"Erro ao exportar PDF: {str(e)}")

    @staticmethod
    def _export_docx(text: str, file_path: str, metrics_original: dict = None, metrics_simplified: dict = None) -> None:
        """
        Exporta texto e métricas para um arquivo DOCX.

        Utiliza a biblioteca python-docx para criar um documento DOCX com o texto e as métricas fornecidos.

        Parâmetros:
            text (str): O texto a ser exportado para o DOCX.
            file_path (str): Caminho onde o arquivo DOCX será salvo.
            metrics_original (dict): Métricas do texto original.
            metrics_simplified (dict): Métricas do texto simplificado.

        Retorna:
            None

        Exceções:
            - Exception: Se ocorrer um erro durante a criação do DOCX.
        """
        try:
            doc = Document()

            # Adicionar título
            doc.add_heading('Texto Simplificado e Traduzido:', level=1)
            doc.add_paragraph(text)

            if metrics_original and metrics_simplified:
                metric_names = {
                    'flesch_reading_ease': 'Índice de Flesch Reading Ease',
                    'flesch_kincaid_grade': 'Grau de Flesch-Kincaid',
                    'smog_index': 'Índice SMOG',
                    'coleman_liau_index': 'Índice de Coleman-Liau',
                    'automated_readability_index': 'Índice ARI',
                    'dale_chall_readability_score': 'Pontuação de Dale-Chall'
                }

                # Métricas do texto original
                doc.add_heading('Métricas do Texto Original:', level=2)
                for key, value in metrics_original.items():
                    metric_name = metric_names.get(key, key)
                    doc.add_paragraph(f"{metric_name}: {value:.2f}")

                # Métricas do texto simplificado
                doc.add_heading('Métricas do Texto Simplificado:', level=2)
                for key, value in metrics_simplified.items():
                    metric_name = metric_names.get(key, key)
                    doc.add_paragraph(f"{metric_name}: {value:.2f}")

            doc.save(file_path)
        except Exception as e:
            raise Exception(f"Erro ao exportar DOCX: {str(e)}")

    @staticmethod
    def _export_txt(text: str, file_path: str, metrics_original: dict = None, metrics_simplified: dict = None) -> None:
        """
        Exporta texto e métricas para um arquivo TXT.

        Abre (ou cria) o arquivo de texto e escreve todos os conteúdos fornecidos, incluindo as métricas.

        Parâmetros:
            text (str): O texto a ser exportado para o TXT.
            file_path (str): Caminho onde o arquivo TXT será salvo.
            metrics_original (dict): Métricas do texto original.
            metrics_simplified (dict): Métricas do texto simplificado.

        Retorna:
            None

        Exceções:
            - Exception: Se ocorrer um erro durante a escrita no TXT.
        """
        try:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write("Texto Simplificado e Traduzido:\n")
                f.write(text)
                f.write("\n\n")

                if metrics_original and metrics_simplified:
                    metric_names = {
                        'flesch_reading_ease': 'Índice de Flesch Reading Ease',
                        'flesch_kincaid_grade': 'Grau de Flesch-Kincaid',
                        'smog_index': 'Índice SMOG',
                        'coleman_liau_index': 'Índice de Coleman-Liau',
                        'automated_readability_index': 'Índice ARI',
                        'dale_chall_readability_score': 'Pontuação de Dale-Chall'
                    }

                    f.write("Métricas do Texto Original:\n")
                    for key, value in metrics_original.items():
                        metric_name = metric_names.get(key, key)
                        f.write(f"{metric_name}: {value:.2f}\n")

                    f.write("\nMétricas do Texto Simplificado:\n")
                    for key, value in metrics_simplified.items():
                        metric_name = metric_names.get(key, key)
                        f.write(f"{metric_name}: {value:.2f}\n")
        except Exception as e:
            raise Exception(f"Erro ao exportar TXT: {str(e)}")
